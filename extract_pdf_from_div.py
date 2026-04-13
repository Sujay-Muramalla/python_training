import sys
from pathlib import Path
from urllib.parse import urlparse, parse_qs, unquote

import requests
from bs4 import BeautifulSoup
import pdfplumber
from docx import Document
from docx.shared import Pt


def extract_pdf_url_from_html(html_content: str) -> str | None:
    """
    Given an HTML snippet (e.g., a <div> containing a PDF viewer iframe),
    find and return the actual PDF URL.

    It handles cases like:
      <iframe src=".../viewer.html?file=https%3A%2F%2F...something.pdf%3FExpires%3D...">
      <iframe src="https://example.com/some.pdf">

    Returns:
      PDF URL as a string, or None if not found.
    """
    soup = BeautifulSoup(html_content, "lxml")

    for tag in soup.find_all(["iframe", "embed"]):
        src = tag.get("src")
        if not src:
            continue

        parsed = urlparse(src)

        # Case 1: src itself points directly to a PDF
        if parsed.path.lower().endswith(".pdf"):
            return src

        # Case 2: PDF URL is inside a query parameter like ?file=<pdf_url>
        qs = parse_qs(parsed.query)
        if "file" in qs and qs["file"]:
            # URL may be percent-encoded
            pdf_url = qs["file"][0]
            pdf_url = unquote(pdf_url)
            return pdf_url

    return None


def download_pdf(pdf_url: str, output_path: Path) -> None:
    """
    Download the PDF from pdf_url to output_path.
    """
    print(f"🔗 Downloading PDF from: {pdf_url}")
    resp = requests.get(pdf_url, stream=True)
    resp.raise_for_status()

    with output_path.open("wb") as f:
        for chunk in resp.iter_content(chunk_size=8192):
            if chunk:
                f.write(chunk)

    print(f"📄 PDF saved to: {output_path}")


def add_paragraph(doc: Document, text: str, style: str | None = None):
    """
    Helper to add a paragraph with consistent spacing.
    """
    text = text.strip()
    if not text:
        return

    p = doc.add_paragraph(text)
    if style:
        p.style = style

    fmt = p.paragraph_format
    fmt.space_before = Pt(4)
    fmt.space_after = Pt(4)


def pdf_to_docx(pdf_path: Path, docx_path: Path):
    """
    Convert a PDF file to a Word document.

    Very simple heuristic:
      - Lines that are ALL CAPS and not too long -> Heading 2
      - Everything else -> Normal
    """
    doc = Document()

    add_paragraph(doc, "Extracted PDF Content", style="Title")
    add_paragraph(
        doc,
        "This document was generated from a PDF found inside an HTML PDF viewer (iframe).",
        style="Normal",
    )

    with pdfplumber.open(str(pdf_path)) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            if not text.strip():
                continue

            lines = text.splitlines()

            # Add a small page marker (optional)
            add_paragraph(doc, f"--- Page {page_num} ---", style="Heading 3")

            for line in lines:
                raw = line.strip()
                if not raw:
                    doc.add_paragraph("")  # blank line
                    continue

                # Heuristic: treat short ALL CAPS lines as headings
                if raw.isupper() and len(raw.split()) <= 8:
                    add_paragraph(doc, raw, style="Heading 2")
                else:
                    add_paragraph(doc, raw, style="Normal")

    doc.save(str(docx_path))
    print(f"✅ Word document created: {docx_path}")


def main():
    if len(sys.argv) < 3:
        print("Usage: python extract_pdf_from_div.py input_div.html output.docx")
        sys.exit(1)

    input_html_path = Path(sys.argv[1])
    output_docx_path = Path(sys.argv[2])

    if not input_html_path.exists():
        print(f"❌ Input file not found: {input_html_path}")
        sys.exit(1)

    # Read the HTML snippet that contains the div with the PDF viewer
    html_content = input_html_path.read_text(encoding="utf-8")

    # 1. Extract the PDF URL from the iframe/embed
    pdf_url = extract_pdf_url_from_html(html_content)
    if not pdf_url:
        print("❌ Could not find a PDF URL inside the provided HTML snippet.")
        sys.exit(1)

    # 2. Download the PDF to a temporary file next to the HTML
    pdf_path = input_html_path.with_suffix(".pdf")
    download_pdf(pdf_url, pdf_path)

    # 3. Convert the downloaded PDF to a Word document
    pdf_to_docx(pdf_path, output_docx_path)


if __name__ == "__main__":
    main()
