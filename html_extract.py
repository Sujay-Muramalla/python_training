import pdfplumber
from docx import Document
from docx.shared import Pt
from pathlib import Path
import sys


def add_paragraph(doc, text, style=None):
    if not text.strip():
        return
    p = doc.add_paragraph(text)
    if style:
        p.style = style
    fmt = p.paragraph_format
    fmt.space_before = Pt(4)
    fmt.space_after = Pt(4)


def pdf_to_docx(pdf_path: str, docx_path: str):
    doc = Document()

    add_paragraph(doc, "Extracted from PDF", style="Title")

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            # Split into lines
            for line in text.splitlines():
                line = line.strip()
                if not line:
                    doc.add_paragraph("")  # blank line
                    continue

                # Very naive heading heuristic: all-caps + short line
                if line.isupper() and len(line.split()) <= 8:
                    add_paragraph(doc, line, style="Heading 2")
                else:
                    add_paragraph(doc, line, style="Normal")

    doc.save(docx_path)
    print(f"✅ Created: {docx_path}")


def main():
    if len(sys.argv) < 3:
        print("Usage: python pdf_to_word.py input.pdf output.docx")
        return

    pdf_path = Path(sys.argv[1])
    docx_path = Path(sys.argv[2])

    if not pdf_path.exists():
        print(f"❌ PDF not found: {pdf_path}")
        return

    pdf_to_docx(str(pdf_path), str(docx_path))


if __name__ == "__main__":
    main()
